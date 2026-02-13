'''
 Telegram Bot for Anketing 
 version 0.1
 Module anketa.py 
  
'''

import io
import re
import logging
import asyncio
import os.path
import sys
import gettext
import json
from datetime import datetime
import requests
from telethon import TelegramClient, events
from telethon.tl.types import  PeerChannel, PeerUser, UpdateNewMessage
from telethon.tl.custom import Button
from telethon import errors
from telethon.events import StopPropagation
from telethon.sessions import StringSession
from apscheduler.schedulers.asyncio import AsyncIOScheduler
import pandas as pd
import filetype
import docx
#from requests.packages.urllib3.util.retry import Retry
# --------------------------------
import settings as sts
import dbmodule as dbm
# --------------------------------
#Glogal vars
bot = None

async def add_admins(event):
    ''' Select users for add to admins list
        event = bot event handled id
        level = user level for show menu exxtended or no
    '''
    id_user = event.query.user_id
    logging.debug(f"Create select users dialog for user {id_user}")
    
    buttons = [
    {
        "text":"👥 Выбор Админа",
        "request_users": {
            "request_id": 1,# button id
            "max_quantity": 5,
            "user_is_bot": False,
            "request_name": True,
            "request_username": True, 
        }
    }
    ]
    reply_markup = {"keyboard": [buttons], "resize_keyboard": True, "one_time_keyboard": True }
    payload = {
    "chat_id": id_user, # Id user to
    "text": f"Нажмите кнопку 'Выбор Админа' чтобы добавть в список администраторов", 
    "reply_markup": json.dumps(reply_markup)
    }

    # Send selection user Button 
    url = f"https://api.telegram.org/bot{sts.mybot_token}/sendMessage"

    response = requests.post(url, data=payload, timeout = 30, proxies=sts.proxies)
    logging.debug(f"Rеsponse Select user button post:{response}\n")

    # hanled answer
    @bot.on(events.Raw(types=UpdateNewMessage))
    async def on_requested_peer_user(event_select):
        logging.debug(f"Get select user event:{event_select}")
        users_id_list=[] 
        usernames=[]
        nicknames=[]
        text_reply=''
        new_admins={}


        try:
            if event_select.message.action.peers[0].__class__.__name__ == "RequestedPeerUser":
                button_id = event_select.message.action.button_id
                if button_id == 1:
                    for peer in event_select.message.action.peers:
                        if peer.user_id in sts.Admins.keys():
                           text_reply=text_reply+f"⚠️{peer.username} {peer.first_name} уже админ!\n"
                           continue
                        #if peer.username in sts.Admins:
                        #   text_reply=text_reply+f"⚠️{peer.username} уже админ!\n"
                        #   continue
                        #if peer.username == None:
                        #   text_reply=text_reply+f"⚠️{peer.first_name} не имеет nickname!\n"
                        #   continue 
                        #usernames.append(peer.first_name)
                        #users_id_list.append(peer.user_id)
                        #nicknames.append(peer.username)
                        new_admins[int(peer.user_id)]=peer.username,peer.first_name

                    bot.remove_event_handler(on_requested_peer_user)
                    if new_admins:
                        logging.debug(f"Get selected users:{new_admins}")
                        # Add new admins in DB
                        async with dbm.DatabaseBot(sts.db_name) as db:
                            ret = await db.db_add_admins(new_admins)
                        if ret:
                            #Update current list of admins
                            #for nik in nicknames:
                            #    sts.Admins.append(nik)
                            sts.Admins.update(new_admins)
                            text_reply=text_reply+f"🏁Администраторы добавлены🏁"
                        else:
                            text_reply=f"🏁Ошибка добавления админа🏁"
                    else:
                        text_reply=text_reply+f"Некого добавить!"

                    reply_markup = { "remove_keyboard": True }
                    payload_remove_kb = {
                    "chat_id": id_user, # Id user to
                    "text": text_reply, 
                    "reply_markup": json.dumps(reply_markup)
                    }
                    response = requests.post(url, data=payload_remove_kb, timeout = 30, proxies=sts.proxies)
                    logging.debug(f"Rsponse Remove keyboard:{response}\n")
                    await create_admin_menu(0, event) 
                    
                    return 
        except Exception as error :
            logging.debug(f"It is not RequestedPeerUser message:{error}")
            return None

async def del_admins(event):
    ''' Delete admins form list
        event = bot event handled id
        level = user level for show menu exxtended or no
    '''
    
    logging.debug("Call del_admins() function")
    bdata_id='DEL_ADMIN_'
    button=[]
    i=0
    admin_name=''
    admin_nickname=''
    logging.debug(f"Len Admins: {len(sts.Admins)}")
   
    if len(sts.Admins) > 1:
        message="❌ Выберете админа для удаления:"        
        for admin_id, cur_admin in sts.Admins.items():
            if i == 0: 
                i=i+1
                continue
            bdata=bdata_id+str(admin_id)
            if cur_admin[1]:
                admin_name = cur_admin[1]
                if cur_admin[0]: admin_nickname = f'@{cur_admin[0]}'
            elif cur_admin[0]: 
                admin_name = cur_admin[0]
            else:
                admin_name ='Noname'

            button.append([ Button.inline(f'👮 {admin_name} {admin_nickname} ({admin_id})', bdata)])
            admin_nickname =''
    else:
           await event.respond("⚠️Нет админов для удаления.")
           await create_admin_menu(0,event)
           return False
    
    await event.respond(message, buttons=button)    
    return True

async def show_admins(event):
    ''' Show current admins
        event = bot event handled id
        level = user level for show menu exxtended or no
    '''
    Builtin_Admin='💂‍♂️'
    Simply_Admin='👮'

    i=True

    admin_name=''
    admin_nickname=''
    rstr='📃Cписок текущих Админов:\n\n'
    for admin_id, cur_admin in sts.Admins.items(): 
        if cur_admin[1]:
            admin_name = cur_admin[1]
            if cur_admin[0]: admin_nickname = f'@{cur_admin[0]}'
        elif cur_admin[0]: 
            admin_name = f'@{cur_admin[0]}'
        else:
            admin_name ='Noname'

        if i:
            rstr=rstr+f'{Builtin_Admin} {admin_name} {admin_nickname} ({admin_id})\n'
            i=False
            admin_nickname = ''
        else:
             rstr=rstr+f'{Simply_Admin} {admin_name} {admin_nickname} ({admin_id})\n'
             admin_nickname = ''

    await event.respond(rstr)
    await create_admin_menu(0, event)

async def check_nickname(username):
    res = {}
    try:
        # Пытаемся получить информацию о пользователе/канале по нику
        entity = await bot.get_entity(username)
        #logging.debug(f"Nickname [{username}] ok. Object type is: {type(entity).__name__}")
        logging.debug(f"User ID: {entity.id}")
        logging.debug(f"First Name: {entity.first_name}")
        logging.debug(f"Username-nickname: {entity.username}")
        res[entity.id]=entity.username,entity.first_name
        return res
    except ValueError:
        # Возникает, если Telethon не нашел сущность (обычно если ника нет)
        logging.debug(f"Nickname [{username}] not found ")
        return False
    except Exception as e:
        logging.debug(f"Error check Nickname [{username}] {e}")
        return False

async def is_utf8_text_file(filename): #NOTUSE now
    """Checks if a file can be entirely decoded as UTF-8 text."""
    try:
        with open(filename, 'r', encoding='utf-8') as file:
            file.read()
        return True
    except UnicodeDecodeError:
        # This exception is raised if the file contains byte sequences 
        # that are invalid for UTF-8 encoding.
        return False
    except Exception as e:
        # Handle other potential exceptions (e.g., file not found, permission errors)
        logging.warning(f"An error occurred: {e}")
        return False

async def get_excel_data(filename, sheet_name=0):
    """
    Reads data from an Excel file into a pandas DataFrame.
    sheet_name can be an integer (0 for the first sheet) or a string ('Sheet1').
    """
    try:
        df = pd.read_excel(filename, sheet_name=sheet_name, header=None )
        # Convert the DataFrame to dict
        res=df.to_dict(orient='split', index=False) 
        return res
    except Exception as e:
        logging.warning(f"Error reading Excel file: {e}")
        return False

async def get_word_text(filename): #NOTUSE now
    """
    Extracts all text from a .docx file.
    """
    document = docx.Document(filename)
    full_text = []
    for paragraph in document.paragraphs:
        full_text.append(paragraph.text)

    # Join paragraphs with a newline character
    return '\n'.join(full_text)

async def get_oldword_text(filename): #FIXME It`s dont work. NOTUSE now
    """
    Extracts all text from old a .doc file.
    """
    # OLD word file - .doc
       # Convert the document and extract text
       
    #text = docx2txt.process(filename) 
    #return text
    pass
    return None
    
async def get_txt_text(filename): #NOTUSE now
    '''
    Get data fron text file 
    '''
    with open(filename, 'r', encoding="utf-8") as file:
                #text = text + [for line in file.readlines()]
                text=file.read()
    
    return text

async def get_new_questions(filename):
    '''
    Docstring для get_new_questions
    Get new questions from file txt,docx,xls,xlsx and return list
    :param filename: file with questions
    '''
        
    root,ext = os.path.splitext(filename)
    
    kind = filetype.guess(filename)
    
    #logging.debug(f'File extension: {kind.extension}')
    #logging.debug(f'File MIME type: {kind.mime}')

    if kind is None:
        logging.debug(f'Cannot guess file type filename: {filename}!')
        return 1
    elif kind.extension == 'xlsx' or kind.extension == 'xls':
        text_content = await get_excel_data(filename)
        logging.debug(f'Xlsx or xls content is:{text_content}')
    
    if not text_content:
        return False
    
    qlist={}
    val=[]
    for item in text_content['data']:
        #item - one question and variants answers if exist
        nan_list=pd.isna(item)
        i=False
        # variants answer to list values dict        
        for x, y in zip(item,nan_list):
            if not y and i:
                val.append(x) 
            i=True
        
        qlist[item[0]]=val
        val=[]

    #qlist = [item.strip() for item in text_content.split('\n')]
    #qlist = list(filter(None, qlist))
    return qlist

async def create_admin_menu(level, event):
    ''' Create Admin menu '''
    logging.debug("Create menu buttons")
    keyboard = [
        [
            Button.inline("📈 Показать статистику", b"/am_stats")
        ],
        [
            Button.inline("📃 Пройти анкетирование", b"/am_anketa")
        ],
        [
            Button.inline("📊 Получить результаты", b"/am_answers")
        ],
        [
            Button.inline("📑 Текущие вопросы", b"/am_show_questions")
        ],
        [
            Button.inline("⬆️ Загрузить новые вопросы", b"/am_questions")
        ]
        ,
        [
            Button.inline("👮‍♂️ Добавть администратора", b"/am_add_admins")
        ]
        ,
        [
            Button.inline("🙅‍♂️ Удалить администратора", b"/am_del_admins")
        ]
        ,
        [
            Button.inline("🕵️ Просмотреть всех админов", b"/am_show_admins")
        ]
    ]
    #clear old message
    await event.delete()
    # send menu
    await event.respond("**☣ Режим Администратора:**", parse_mode='md', buttons=keyboard)

async def show_stats(event):
    '''
    show statistics for users
    '''
    logging.debug("Call show_stats() function")

    async with dbm.DatabaseBot(sts.db_name) as db:
        rows = await db.get_info_by_users()
    if not rows:
        await event.respond(f"🚷На данный момент нет информаци.\nЕще никто не прошел опрос.")
        return False

    strstat=f"🔢 Ответили на вопросы: {len(rows)}\n\n👥 Список прошедших опрос:\n\n"

    for row in rows:
        #dt = datetime.strptime(dict(row).get('date'),'%Y-%m-%d %H:%M:%S.%f')
        #strstat=strstat+f"{dict(row).get('name_user')} { dt.strftime('%d.%m.%y %H:%M') }\n"
        strstat=strstat+f"{dict(row).get('name_user')}\n"

    await event.respond(strstat)
  
    return True 

async def send_answ_db(event):
    '''
    send Answers DB to Admin (load results)
    '''
    logging.debug("Call send_answ_db() function")

    dt = datetime.now().strftime('%d%m%Y_%H%M%S')
    
    filename = f"reports/report_{dt}.xlsx"
    logging.debug(f"Gen filename: {filename}")
    res = await gen_excel(filename)
    if res:
        message="Ваш отчет"
        await bot.send_file( event.query.user_id, filename, caption=message, parse_mode="html" ) 
        return True
    else:
        await event.respond(f"🚷На данный момент нет информаци для отчета.\nЕще никто не прошел опрос.")
        return False

async def gen_excel(filename):
    '''
    Generate excel table
    '''
    data={}
    data['name_user']=[]
    data['nick_user']=[]
    data['question_id']=[]
    data['answer_user']=[]
    data['date']=[]
    data['time']=[]

    
    async with dbm.DatabaseBot(sts.db_name) as db:
        rows = await db.get_info_for_report()
    if not rows:
        return False

    # Get name_user, nick_user, question_id, answer_user, date
    for row in rows:
        data['name_user'].append(dict(row).get('name_user'))       
        data['nick_user'].append(dict(row).get('nick_user'))
        index=int(dict(row).get('question_id'))
        #data['question_id'].append(all_questions[index-1])
        key_q=list(all_questions)[index-1]
        data['question_id'].append(key_q)
        answer_cur=dict(row).get('answer_user')
        logging.info(f"Results gen excel: answer_cur:{answer_cur} all_questions.get(key_q):{all_questions.get(key_q)}")
        if all_questions.get(key_q):
            data['answer_user'].append(all_questions.get(key_q)[int(answer_cur)-1])
        else: 
            data['answer_user'].append(answer_cur)
        #2024-03-03 11:46:05.488155
        dt = datetime.strptime(dict(row).get('date'),'%Y-%m-%d %H:%M:%S.%f')
        date = dt.strftime('%d.%m.%y')
        time = dt.strftime('%H:%M')
        data['date'].append(date)
        data['time'].append(time)

    df = pd.DataFrame(data)

    # Order the columns if necessary.
    #df = df[["Rank", "Country", "Population"]]

    # Create a Pandas Excel writer using XlsxWriter as the engine.
    writer = pd.ExcelWriter(filename, engine="xlsxwriter")

    # Write the dataframe data to XlsxWriter. Turn off the default header and
    # index and skip one row to allow us to insert a user defined header.
    df.to_excel(writer, sheet_name="Общий", startrow=1, header=False, index=False)

    # Get the xlsxwriter workbook and worksheet objects.
    workbook = writer.book
    worksheet = writer.sheets["Общий"]
    #worksheet = writer.sheets["Sheet1"]

    # Get the dimensions of the dataframe.
    (max_row, max_col) = df.shape

    # Create a list of column headers, to use in add_table().
    column_settings = [{"header": column} for column in df.columns]

    # Add the Excel table structure. Pandas will add the data.
    worksheet.add_table(0, 0, max_row, max_col - 1, {"columns": column_settings})

    # Make the columns wider for clarity.
    worksheet.set_column(0, max_col - 1, 12)

    # Close the Pandas Excel writer and output the Excel file.
    writer.close()
    
    return True
    
async def get_qusetion_data(event_bot):
    '''
    get and load questions to DB Questions
    '''
    logging.debug("Call get_qusetion_data() function")
    
    await event_bot.respond(\
    "📎 Загрузите файл с вопросами.\n\n" \
    "Поддержиаются следующие типы файлов:\n" \
    #"🔹Текстовый файл (txt) по одному вопросу на строке\n" \
    #"🔹MS Word файл (docx) по одному вопросу на строке\n" \
    "🔹MS Excel файл (xls,xlsx) по одному вопросу в ячейке в первой колонке\n" \
    "🔹варианты ответов в следующих за вопросом колонках\n" \
    "🔹если нет варианта ответа - ответ вводит опрашиваемый\n" \
    #"⚠️ Старый формат MS word (doc) не поддерживается!\n" \
    "\n♨️ Текущие вопросы и ответы будут удалены!")

    @bot.on(events.NewMessage())
    async def bot_handler_f_bot(event):
        #logging.debug(f"Get NewMessage event_bot: {event}")      
        if event.message.document:
            download_path = await event.message.download_media(file="questionfiles/") 
            logging.info(f'File saved to: {download_path}')                                   
            #with open(download_path, 'r', encoding="utf-8") as file:
            #    new_questions = [line.strip() for line in file.readlines()]
            new_questions = await get_new_questions(download_path)
            if not new_questions:
                await event_bot.respond("⚠️Данный формат файла не поддерживается!")
                bot.remove_event_handler(bot_handler_f_bot)
                await create_admin_menu(0, event_bot)
                return False
            all_questions.clear()   
            all_questions.update(new_questions)
            logging.info(f'New all_questions: {all_questions}')
            async with dbm.DatabaseBot(sts.db_name) as db:
                await db.db_rewrite_new_questions(all_questions)

            await event.respond("Данные загружены в бот.")
            bot.remove_event_handler(bot_handler_f_bot)
            await create_admin_menu(0, event_bot)
    
async def home():
    '''
    stub function
    '''
    logging.debug("Call home stub function")
    return 0

async def check_user_run_anketa(id_user, event_bot, menu):
    '''
    Test user already answer or not
    and continue
    '''    
    async with dbm.DatabaseBot(sts.db_name) as db:
        res = await db.db_exist_id_user(id_user)
    
    logging.info(f"Exist_id_user: {res}")

    # if user already answer     
    if res:
       #await event_bot.respond(f"Вы уже отвечали на вопросы.\n Желаете пройти опрос снова?\n Предыдущие ответы будут потяряны.\n")
       keyboard = [ Button.inline("Да", b"/yes"),Button.inline("Нет", b"/no") ]
       await event_bot.respond("⚠️Вы уже отвечали на вопросы.\nЖелаете пройти опрос снова?\nПредыдущие ответы будут потеряны.\n", parse_mode='md', buttons=keyboard)
      
       @bot.on(events.CallbackQuery())
       async def callback_yn(event):            
            button_data = event.data.decode()
            logging.info(f"Callback yes/no: {button_data}")
            #await event.delete()
            if button_data == '/no':
                await event_bot.respond(f"До свидания.\n\n")
                bot.remove_event_handler(callback_yn)                
            elif button_data == '/yes': 
                async with dbm.DatabaseBot(sts.db_name) as db:
                    row = await db.db_del_user_answers(id_user)
                bot.remove_event_handler(callback_yn)
                await run_anketa(id_user, event_bot, menu)                                      
            return 0
    else:
        await run_anketa(id_user, event_bot, menu)       
        return 2

async def run_anketa(id_user, event_bot, menu):
    '''
    run main process for anketting
    '''
    user_ent = await bot.get_entity(id_user)
    sender = await event_bot.get_sender()
    sender_id = sender.id
    nickname = user_ent.username
    first_name = user_ent.first_name
    question_id=0
    button=[]
    bdata=''
    v=1
    answ_v=[]
    answers={}
    await event_bot.respond(f"Ответе пожалуйста на несколько вопросов\n"\
                            "⚠️На каждый ответ отводится {sts.TIMEOUT_FOR_ANSWER} секунд.\n\n")

    async with bot.conversation(id_user) as conv:
        def my_press_event(id_user):
            return events.CallbackQuery(func=lambda e: e.sender_id == id_user) #FIXME Need or not use pattern for get button?
        try:
            for cur_question  in all_questions:
                if not all_questions.get(cur_question):
                    await conv.send_message(f"Вопрос {question_id+1}:\n{cur_question}")
                    response = await conv.get_response(timeout=sts.TIMEOUT_FOR_ANSWER)
                    resp_text = response.text
                    logging.info(f"Get respond text: {question_id} : {resp_text}")
                    answers[question_id+1]=resp_text
                else:
                    button.clear()
                    str_qst=f"Вопрос {question_id+1}:\n{cur_question}"
                    v=1
                    for variant in all_questions.get(cur_question):
                        bdata=f'VARIANT_{question_id}_{v}'
                        button.append([ Button.inline(f'🔹 {variant}', bdata)])
                        v=v+1
                    await conv.send_message(str_qst, buttons=button)
                    handle = conv.wait_event(my_press_event(sender_id),timeout=sts.TIMEOUT_FOR_ANSWER) #FIXME Need or not use pattern for get button?
                    event_res = await handle
                    button_pressed = event_res.data.decode('utf-8')
                    answ_v = button_pressed.replace('VARIANT_', '').split('_')
                    logging.debug(f"Get respond button text: {question_id} : {button_pressed} : {answ_v}")
                    answers[question_id+1]=answ_v[1]
                question_id = question_id + 1
            logging.debug(f"Dict All answers: {answers}")
            # Write Answers to DB
            async with dbm.DatabaseBot(sts.db_name) as db:     
                    await db.db_add_answer(id_user, first_name, nickname, answers)
            await conv.send_message(f"🔆 Вы ответили на все вопросы.\n"\
                                    "Результаты сохранены.\n"
                                    "Для повторного прохождения опроса\n"\
                                    "нажмите кнопку Старт\n")
        except TimeoutError as error:
            logging.debug(f"Get timeout {sts.TIMEOUT_FOR_ANSWER} sec for user {id_user} on answer {cur_question} ")
            await conv.send_message(f"⚠️Отведенное время {sts.TIMEOUT_FOR_ANSWER} секунд на ответ истекло.\n"\
                                    "Результаты не будут сохранены.\n"\
                                    "Пожалуйста пройдите опрос заново.\n"\
                                    "Для этого нажмите кнопку Старт\n")
        conv.cancel()
        if menu: 
            await create_admin_menu(menu, event_bot)

    return 0 

async def show_qusetions(event_bot):
    '''
    Show all questions
    '''
    i=1
    message=f"🧐 Текущие вопросы:"
    for qst in all_questions:
        message = message + f"\n{i}.{qst}\n"
        for variant in all_questions.get(qst):
           message = message + f"  🔹 {variant}\n" 
        i=i+1
    
    await event_bot.respond(message)
    await create_admin_menu(0, event_bot)

async def main_frontend():
    ''' Loop for bot connection '''
    
    #global all_questions

    @bot.on(events.NewMessage())
    async def bot_handler_nm_bot(event_bot):
        logging.debug(f"Get NewMessage event_bot: {event_bot}")
        menu_level = 0
      
        id_user = event_bot.message.peer_id.user_id
        logging.info(f"LOGIN USER_ID:{id_user}")
        #user_ent = await bot.get_entity(id_user)
        #nickname = user_ent.username
        #first_name = user_ent.first_name
        
        #logging.debug(f"Get username for id {id_user}: {nickname}")


        if event_bot.message.message == '/start':
            if id_user in sts.Admins.keys():
                #await event_bot.respond("You are admin!")
                await create_admin_menu(0, event_bot)
            else:
                # run anketa for all users who not Admin                    
                await check_user_run_anketa(id_user, event_bot, 0)           
        #elif event_bot.message.message == '/am_stats' and permissions.is_admin:
        #    await show_stats(event_bot)
        #    await create_admin_menu(0, event_bot)
        #elif event_bot.message.message == '/am_anketa'  and permissions.is_admin:
        #    await check_user_run_anketa(id_user, event_bot, 1)
        #    await create_admin_menu(0, event_bot)
        #elif event_bot.message.message == '/am_answers'  and permissions.is_admin:
        #    await send_answ_db(event_bot)
        #    await create_admin_menu(0, event_bot)
        #elif event_bot.message.message == '/am_questions' and permissions.is_admin:
        #    all_questions = await get_qusetion_data(event_bot)
        #    await create_admin_menu(0, event_bot)
        #else:     
        #    pass

    # Run hundler for button callback - menu for Admin
    @bot.on(events.CallbackQuery())
    async def callback_bot_choice(event_bot_choice):
        id_user = event_bot_choice.query.user_id
        user_ent = await bot.get_entity(id_user)
        nickname = user_ent.username
        logging.debug(f"Get callback event for user[{id_user}] {event_bot_choice}")
       
        # If user not Admin ignore button actions  
        if id_user not in sts.Admins.keys(): return 0

        button_data = event_bot_choice.data.decode()
        #await event_bot.delete()
        if button_data == '/am_stats':
            await show_stats(event_bot_choice)
            await create_admin_menu(0, event_bot_choice)
        elif button_data == '/am_anketa':
            await check_user_run_anketa(id_user, event_bot_choice, 1)
            #await create_admin_menu(0, event_bot_choice)
        elif button_data == '/am_answers':
            await send_answ_db(event_bot_choice)
            await create_admin_menu(0, event_bot_choice)
        elif button_data == '/am_questions':
            await get_qusetion_data(event_bot_choice)
            #await create_admin_menu(0, event_bot_choice)
        elif button_data == '/am_show_questions':
            await show_qusetions(event_bot_choice)
            #await create_admin_menu(0, event_bot_choice)
        elif button_data == '/am_add_admins':
            await add_admins(event_bot_choice)
        elif button_data == '/am_del_admins':
            await del_admins(event_bot_choice)
        elif button_data == '/am_show_admins':
            await show_admins(event_bot_choice)
        elif  'DEL_ADMIN_' in button_data:
            # Delete admin
            data = button_data
            admin_id_delete = int(data.replace('DEL_ADMIN_', ''))
            async with dbm.DatabaseBot(sts.db_name) as db:
                res = await db.db_del_admins(admin_id_delete)
            logging.info(f'All:{sts.Admins} admin_id_delete:_{admin_id_delete}_')
            sts.Admins.pop(admin_id_delete)
            await event_bot_choice.respond(f"🏁Админ {admin_id_delete} удален🏁")
            await create_admin_menu(0, event_bot_choice)
    return bot

async def main():
    ''' Main function '''

    print("Start anketa Bot...")
    
    # Check for Admin and get user_id, clear and create new dict Admins for 
    # full data about Admin
   
    ret = await check_nickname(sts.Builtin_admin)
    if not ret:
        logging.error(f'Admin with nickname: {sts.Builtin_admin} Not exist in Telegram! Check config file!')
        print(f'Admin with nickname: {sts.Builtin_admin} Not exist in Telegram! Check config file!')
        exit(-1)
    else:
        sts.Admins.clear()
        sts.Admins.update(ret)
        #print(f'Admin with nickname: {sts.Admins}')
        #sexit(-1)


    async with dbm.DatabaseBot(sts.db_name) as db:
        logging.debug('Create db if not exist.')
        await db.db_create()
        new_questions = await db.db_load_questions()
        rows = await db.db_load_admins()
        adm = {}
        if rows:
            for row in rows:
                #sts.Admins.append(dict(row).get('admin'))
                adm[dict(row).get('admin_id')]=dict(row).get('admin_nickname'),dict(row).get('admin_firstname')
        sts.Admins.update(adm)

        logging.info(f"Get Admins from db: {adm}\n")
        logging.info(f'All:{sts.Admins}\n')

    if new_questions:
        all_questions.clear()
        all_questions.update(new_questions)

    # Run basic events loop
    await main_frontend()    

#------------------- Main begin -----------------------------------------------

sts.get_config()
# Enable logging

# Init default questions
all_questions = {   "text_q1":[],
                    "text_q2":['variant1','variant2'],
                    "text_q3":['variant1'],
                    "text_q4":['variant1','variant2','variant3'],
                    "text_q5":[]
                }

filename=os.path.join(os.path.dirname(sts.logfile),os.path.basename(sts.logfile))
logging.basicConfig(level=sts.log_level, filename=filename, filemode="a", format="%(asctime)s %(levelname)s %(message)s")
logging.info("Start frontend bot.")

localedir = os.path.join(os.path.dirname(os.path.realpath(os.path.normpath(sys.argv[0]))), 'locales')

if sts.use_proxy:
    prx = re.search('(^.*)://(.*):(.*$)', sts.proxies.get('http'))
    proxy = (prx.group(1), prx.group(2), int(prx.group(3)))
else: 
    proxy = None

# Set type session: file or env string
if not sts.ses_bot_str:
    session = sts.session_bot
    logging.info("Use File session mode.")
else:
    session = StringSession(sts.ses_bot_str)
    logging.info("Use String session mode.")
    
# Init and start Telegram client as bot
bot = TelegramClient(session, sts.api_id, sts.api_hash, system_version=sts.system_version, proxy=proxy).start(bot_token=sts.mybot_token)

#bot.start()

with bot:
    bot.loop.run_until_complete(main())
    bot.run_until_disconnected()


